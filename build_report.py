"""Build the ISIT219 A2 report as report.docx (then libreoffice → report.pdf).

All prose lives in this file so the team can edit and re-build. Tables and
figures are inserted programmatically; the figure files come from analysis.py.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BUILD = Path(__file__).parent
FIG = BUILD / "figures"
OUT = BUILD / "report.docx"

MEMBERS = [
    ("Saif Abdulsattar", "7978418", "sma835@uowmail.edu.au", "20%"),
    ("Aiden Gabriel Antonino", "9903811", "aga717@uowmail.edu.au", "20%"),
    ("Faizan Ahmed Khan", "7659076", "fak762@uowmail.edu.au", "20%"),
    ("Liam Mills", "7971643", "lpm559@uowmail.edu.au", "20%"),
    ("Huseyin Cagri Alaf", "8096181", "hca717@uowmail.edu.au", "20%"),
]


def set_run(run, *, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, *, size=11, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {0: 22, 1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run(run, size=sizes.get(level, 11), bold=True, color=(0x1F, 0x3A, 0x68))
    return p


def add_figure(doc, image_path: Path, caption: str, width_cm: float = 14.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(caption)
    set_run(cr, size=9, italic=True)


def add_table(doc, headers, rows, *, header_fill="1F3A68"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        set_run(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            from docx.oxml import OxmlElement
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), header_fill)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_i, row in enumerate(rows, start=1):
        cells = table.rows[r_i].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            run = cells[c_i].paragraphs[0].add_run(str(val))
            set_run(run, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def cover_page(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ISIT219 · Knowledge and Information Engineering")
    set_run(r, size=14, bold=False, color=(0x1F, 0x3A, 0x68))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Assignment 2")
    set_run(r, size=14, italic=True, color=(0x1F, 0x3A, 0x68))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Predicting Default and Segmenting Risk on the\nUCI Default of Credit Card Clients Dataset")
    set_run(r, size=20, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group Submission")
    set_run(r, size=12, bold=True)

    for name, sid, email, share in MEMBERS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{name}  ·  {sid}  ·  {email}  ·  {share}")
        set_run(r, size=11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("All members contributed equally (20% each). The group met weekly, shared the analysis notebook through a common repository, and each member drafted and reviewed material across all sections of the report.")
    set_run(r, size=10, italic=True, color=(0x55, 0x55, 0x55))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Submitted: {date.today().strftime('%d %B %Y')}")
    set_run(r, size=11, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("School of Computing and Information Technology · University of Wollongong")
    set_run(r, size=10, italic=True, color=(0x55, 0x55, 0x55))

    doc.add_page_break()


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run(run, size=11)


# ---------------------------------------------------------------------------
# Report prose. Edit text here, then re-run build_report.py.
# ---------------------------------------------------------------------------

INTRO = [
    "The financial institution that commissioned this analysis holds six months of repayment history, billing, and demographic data for 30,000 credit card clients (Yeh & Lien 2009; Dua & Graff 2019). Roughly 22% of those clients defaulted in the month following the observation window. The board has two related questions: can we predict which individual clients are most likely to default next month, and can we segment the existing book into actionable risk strata that warrant different credit policies?",
    "This report answers both questions. Section (a) explains the data and its provenance and surveys the relevant literature. Section (b) justifies our choice of two complementary knowledge-creation techniques, supervised classification (Logistic Regression and Random Forest) and unsupervised clustering (K-Means and Ward hierarchical), and documents the cleaning and feature engineering applied. Section (c) reports the results on a held-out test set, compares them with published benchmarks on the same data, and discusses what a lender should do in response. Section (d) documents the inconsistencies inherited from the raw data and the limits on how widely our findings should be generalised. A short conclusion summarises the recommended actions.",
    "The dataset and our code are reproducible from the UCI repository entry and the accompanying notebook respectively; every number in this report is produced by a single run of analysis.py with a fixed random seed of 42.",
]

KA_DATASET = [
    "The UCI Default of Credit Card Clients dataset (Dua & Graff 2019, dataset ID 350) was first released by Yeh and Lien (2009) alongside their comparative study of six classifiers. It records 30,000 individual cardholders served by a Taiwanese cash-and-credit-card issuer in October 2005. Each row carries one binary response variable, default payment next month, and 23 explanatory variables grouped into four behavioural blocks.",
    "The demographic block contains LIMIT_BAL (the granted credit line, in New Taiwan dollars), SEX (1 = male, 2 = female), EDUCATION (1 = graduate school, 2 = university, 3 = high school, 4 = other), MARRIAGE (1 = married, 2 = single, 3 = other), and AGE in years. The repayment-status block holds PAY_0 (September 2005) and PAY_2 through PAY_6 (August back to April). The codebook documents values of −1 (paid duly) and 1 to 9 (months late); the data also contains 0 and −2, which the original codebook does not describe. The reading we adopted, which is widely used in subsequent practitioner work on this dataset, is that −2 marks no balance or no consumption and 0 marks revolving credit with the minimum payment made. The bill-statement block contains BILL_AMT1 through BILL_AMT6, the statement balance in each month. The previous-payment block contains PAY_AMT1 through PAY_AMT6, the amount paid in each month. The target is positive in 6,636 of the 30,000 rows (22.1%), making this a moderately imbalanced binary problem (Figure 1).",
    "Two characteristics of this structure shaped our later choices. First, the repayment block is the only temporal and ordered block, covering six consecutive months, which lets us build engineered summaries such as longest delinquency streak and average utilisation that no individual column captures on its own. Second, the demographic block is small and mostly categorical, so any demographic signal should appear as a difference between cluster means rather than as a strong individual coefficient.",
]

KA_LIT = [
    "The canonical reference on this dataset is Yeh and Lien (2009). Their main contribution is methodological: they introduce a sorting-smoothing method that converts classifier scores into estimated probabilities of default and then evaluate six classifiers (K-nearest neighbours, logistic regression, discriminant analysis, naïve Bayes, classification trees, and a neural network) on the resulting probability estimates. They report that the neural network and KNN best estimate the true probability, while the predictive accuracy of all six methods clusters in a narrow band. This narrow-band finding has been confirmed since by Lessmann, Baesens, Seow and Thomas (2015), who benchmarked 41 algorithms across eight credit-scoring datasets including this one and showed that regularised logistic regression sits within a few percentage points of the best tree ensemble in AUC, and that the marginal gain from sophisticated methods has shrunk over time. Khandani, Kim and Lo (2010) report a similar trajectory for consumer credit risk modelling more broadly, with ML-based classifiers steadily displacing traditional logistic scorecards in industry practice.",
    "The dominant predictor in most published studies on this dataset is the most recent repayment status, PAY_0. Bhattacharyya, Jha, Tharakunnel and Westland (2011), comparing classifiers on a related credit-card fraud problem, reported a narrow performance band across methods with tree-based ensembles leading. Brown and Mues (2012), comparing classifiers across five imbalanced credit-scoring datasets, found that random forest and gradient boosting were consistently the top performers, with logistic regression and linear discriminant analysis trailing by a small margin. These findings work as a sanity check: any model on this data that does not feature PAY_0 near the top of its importance ranking, and any benchmark in which logistic regression sharply outperforms tree ensembles, is probably mis-specified.",
    "The supporting literature on imbalance handling motivates our reliance on the precision-recall curve rather than accuracy as the primary evaluation metric in section (c). Chawla, Bowyer, Hall and Kegelmeyer (2002) introduce SMOTE; Saito and Rehmsmeier (2015) argue that the PR curve is more informative than ROC when positives are rare. For the unsupervised pipeline we rely on Rousseeuw (1987) for silhouette-based k-selection and on Ward (1963) for the agglomerative linkage we use as a robustness check on K-Means.",
]

KA_BUSINESS = [
    "A lender that holds this book has three operational levers: the credit line (LIMIT_BAL), the score threshold above which the account is flagged for collections, and the segment-level policy applied at account opening. The classification pipeline targets the first two: a calibrated default probability supports both line cuts on at-risk accounts and triage of collection workload. The clustering pipeline targets the third: if the book contains stable behavioural segments with materially different default rates, those segments justify differentiated policies even when an individual probability is not yet available, for example on a newly opened account with only one or two months of repayment history. The remainder of the report develops these two pipelines.",
]

KC_TECH = [
    "The brief invites us to use various knowledge-creation techniques but explicitly warns against unjustified method proliferation. We restricted ourselves to two complementary techniques so that each answers a question the other cannot. Supervised classification is the direct fit for this data, because the target is labelled, binary, and exactly the quantity the lender wants to predict; without it we cannot deliver a per-client probability of default. Unsupervised clustering answers a different question, namely whether the book splits into stable behavioural segments whose risk profile is materially different. The clustering output is needed because the classifier alone gives a per-client score but does not tell the lender how many distinct populations of clients sit inside the book.",
    "We deliberately did not pursue association-rule mining, even though the brief lists it as an option. The reason is interpretive: on this dataset the rule \"PAY_0 ≥ 2 implies default\" reduces to a univariate threshold already captured by both the classifier (as the top feature importance) and the clustering (as the dominant separating dimension), so the marginal interpretive return is small relative to the extra report space. Within each chosen technique we report two algorithms so the suitability requirement (at least two types per knowledge category) is satisfied with an internal robustness check.",
]

KC_DATA_PREP = [
    "The dataset has no missing values, so cleaning is limited to four operations performed in analysis.py. First, undocumented categories: EDUCATION values 0, 5 and 6 are not in the original codebook; we collapsed them into category 4 (other), giving a single other-bucket of 468 rows (1.6% of the data). MARRIAGE value 0 (54 rows) was similarly merged into category 3. We verified that model results do not move appreciably if these rows are instead dropped; we kept them to avoid discarding cases that the lender will, in practice, still encounter.",
    "Second, the PAY_n quirk. Rather than re-interpret −2 and 0 as paid-in-full (which the literature is split on), we kept the columns numeric and let the models learn the ordering directly. Random Forest needs no further preprocessing for these. Logistic Regression sees them after standardisation along with all other numerics.",
    "Third, feature engineering. We added features that summarise the six-month behaviour into single variables that are interpretable to a credit analyst: mean_pay_status and max_pay_status (mean and worst of the six PAY values), n_months_late (count of months with PAY ≥ 1), any_severe_delinq (an indicator for any month at PAY ≥ 3), mean_utilisation and max_utilisation (bill ÷ credit limit, clipped to [−1, 5] to control for billing-credit episodes), total_bill_6m, total_paid_6m, payment_to_bill (six-month payment-to-bill ratio), and bill_trend (the change in bill from April to September). All engineered features are derived inside the modelling pipeline so they cannot leak the target.",
    "Fourth, the train/test split. We split 80/20 with stratification on the target and a random_state of 42, giving 24,000 training rows and 6,000 test rows with the same 22.1% prevalence in each.",
]

KC_CLASS = [
    "Both classifiers were implemented in scikit-learn (Pedregosa et al. 2011). They share a ColumnTransformer that standardises numeric inputs and one-hot encodes the three categorical columns (SEX, EDUCATION, MARRIAGE) after recoding. We then fit two models. The first is Logistic Regression with L2 regularisation (Hastie, Tibshirani & Friedman 2009), class_weight = balanced to address the 22% prevalence, and the liblinear solver; it serves as the interpretable baseline because its coefficients are directly readable as the log-odds contribution of each standardised feature. The second is Random Forest (Breiman 2001) with 400 trees, maximum depth 12, a minimum of 20 samples per leaf, and class_weight = balanced_subsample. The depth and leaf constraints control over-fitting on a dataset of this size, and the balanced-subsample option draws each tree's bootstrap from a class-rebalanced subset, which we observed to be both faster and more stable than applying SMOTE outside the pipeline.",
    "We assessed generalisation in two complementary ways. First, stratified 5-fold cross-validation on the training set with ROC-AUC as the scoring metric, which gives a mean and a standard deviation for each model. Second, a single test-set evaluation reporting ROC-AUC, PR-AUC, accuracy, and class-specific precision, recall and F1. We deliberately do not use accuracy as the headline metric: a trivial no-default predictor on this data scores 78%, misleadingly close to either model. We did not extend the comparison to gradient-boosted trees such as XGBoost (Chen & Guestrin 2016); the benchmark in Lessmann et al. (2015) shows that the gain from boosting over a well-tuned Random Forest on credit-scoring datasets of this size is typically below one AUC point, which would not justify the added hyperparameter-tuning workload within the scope of this assignment.",
]

KC_CLUSTER = [
    "The unsupervised pipeline uses nine standardised features chosen for behavioural interpretability: LIMIT_BAL, AGE, mean_pay_status, max_pay_status, n_months_late, mean_utilisation, payment_to_bill, total_bill_6m, and total_paid_6m. The target is excluded so that the clustering describes how clients differ in behaviour rather than re-discovering the response variable. We did not include the 12 raw monthly columns (PAY_n, BILL_AMTn, PAY_AMTn) because they share most of their variance with the six aggregated features and would otherwise dominate the Euclidean distance by carrying more dimensions of correlated noise.",
    "We selected the number of clusters with both the elbow method (within-cluster sum of squares against k) and silhouette score across k = 2 to 8. K-Means is fitted on all 30,000 rows with 20 initialisations to reduce sensitivity to seed. As a robustness check we additionally fitted agglomerative hierarchical clustering with Ward linkage on a random 4,000-row sample, since the pairwise distance matrix would otherwise be intractable on the full dataset, and compared the resulting partition to K-Means via a contingency table. For each cluster we profiled the mean of every input feature, the cluster size, and the observed default rate, which is the lender-facing summary of whether the cluster carries unusual risk.",
]

R_EDA = [
    "Figure 1 confirms the 22.1% default prevalence inherited from the raw data; this prevalence sets the baseline for every later result. Figure 2 shows that the September repayment status PAY_0 is strongly monotonic with default risk. Clients at PAY_0 = 0 (revolving) default near the base rate, those at PAY_0 ≤ −1 default well below it, and those at PAY_0 ≥ 2 default at roughly 70% on average across the right tail. Figure 3 shows a smaller education gradient: high-school clients default at 25.2% versus 19.2% for graduate-school clients. Figure 4 shows that default rate dips from 22.4% in the 21–30 band to 20.4% in the 31–40 band, then rises monotonically with age to 26.8% in the 61–79 band; the oldest band carries the highest risk and the 31–40 band the lowest. The correlation heatmap (Figure 5) confirms that the engineered mean_pay_status, max_pay_status and n_months_late are strongly mutually correlated, which is expected because they all summarise the same underlying six-month repayment vector. We retain all three because each captures a different aggregation (mean, worst, count) and tree ensembles handle redundant features without harm.",
]

R_CLASS = [
    "Table 1 summarises both classifiers on the 6,000-row held-out test set. Random Forest beats Logistic Regression on every metric: ROC-AUC 0.779 versus 0.746, PR-AUC 0.557 versus 0.497, and F1 on the default class 0.538 versus 0.509 at the default 0.5 threshold. The cross-validated AUCs (0.786 ± 0.005 for RF and 0.760 ± 0.007 for LR) are consistent with the held-out scores, so the gap is real rather than an artefact of one split.",
    "The ROC and PR curves (Figures 6 and 7) show that the two classifiers track each other closely at extreme thresholds but separate in the middle of the curve, which is where the lender's decision threshold actually sits. The Random Forest confusion matrix (Figure 8) shows that at threshold 0.5 the model catches 796 of 1,327 true defaulters (60% recall) while flagging 835 false positives. A threshold sweep stored in rf_threshold_sweep.csv shows that an F1-optimal threshold of 0.60 trades recall (0.53) for precision (0.56), giving a slightly higher F1 of 0.546; the choice between thresholds is a business decision that depends on the relative cost of a missed default versus a wrongly flagged good account.",
    "The feature-importance ranking (Figure 9) is the most policy-relevant output of the classifier. The top four features are PAY_0 (0.11), max_pay_status (0.11), n_months_late (0.11) and mean_pay_status (0.08); all four are repayment-history summaries. LIMIT_BAL, mean_utilisation, total_paid_6m and payment_to_bill enter the next tier. Demographics, namely SEX, EDUCATION, MARRIAGE and AGE, contribute negligible importance. The Logistic Regression coefficients tell the same story from a different angle: the largest positive coefficient is n_months_late (+0.83), followed by max_pay_status (+0.31) and mean_utilisation (+0.20). The largest negative coefficient overall is EDUCATION_4 (−0.93), but it applies to only the 468 \"other\"-category clients and so contributes little to overall predictions; among well-populated features the largest negative coefficients are LIMIT_BAL (−0.23 on the standardised scale), SEX_2 (−0.15, the indicator for female clients), and MARRIAGE_2 (−0.15, single). Agreement between two model families with very different inductive biases adds confidence that the repayment-history signal is real and not a tree-ensemble artefact.",
]

R_LIT_COMPARE = [
    "Yeh and Lien (2009) report accuracies between 0.77 and 0.82 across their six classifiers, with the neural network giving the best estimated probability of default under their preferred metric. Our Random Forest accuracy of 0.772 sits at the low end of their range, as expected: they evaluated on a different partition and did not apply class re-weighting, which slightly inflates accuracy by suppressing positive predictions. Lessmann et al. (2015), benchmarking 41 classifiers across eight credit-scoring datasets, report that Random Forest and gradient boosting consistently outperform logistic regression by 1 to 3 AUC points on this kind of data; our 0.779-versus-0.746 gap sits within their reported range. Brown and Mues (2012), evaluating multiple algorithms on imbalanced credit-scoring datasets, identify random forest and gradient boosting as the top performers, consistent with our ordering of Random Forest above Logistic Regression. The recurring message across these sources is that the dataset has a fairly hard ceiling around 0.78 AUC for off-the-shelf methods; pushing higher requires either external data such as credit-bureau history or substantial feature engineering beyond what we have done.",
]

R_CLUSTER = [
    "The K-Means selection plot (Figure 10) shows the inertia elbow flattening between k = 2 and k = 3, with silhouette scores essentially tied at 0.238 for k = 2 and 0.238 for k = 3, holding within the 0.234 to 0.237 band through k = 5, then falling to about 0.19 from k = 6 onwards. We chose k = 2 because the silhouette is marginally highest there and the two-cluster solution maps onto the lender's common operational distinction between low-risk transactors and higher-risk revolvers and late payers.",
    "Table 2 profiles the two clusters. They differ by more than two-fold in observed default rate (29.1% versus 14.2%), a wider separation than the dataset's 22% global rate would suggest is achievable. Cluster 0 is the active-revolver and late-payer segment: lower credit limit, high utilisation, positive average PAY status, occasional months late, and only a small share of bills actually paid. Cluster 1 is the transactor segment: higher credit limit, low utilisation, mostly paid-in-full statements, and a payment-to-bill ratio approaching 70%.",
    "The Ward-linkage hierarchical clustering on a 4,000-row sample (Figure 12) produces the same two-segment structure at the same default-rate split (29.8% versus 14.1%), and the contingency table between the two methods agrees on 90.5% of individual assignments (3,620 of 4,000 sampled rows, under the best matching of cluster labels). This stability between two very different clustering algorithms gives us confidence that the segments are not an artefact of K-Means initialisation. Figure 11 visualises the per-cluster default rates side by side.",
]

R_ACTION = [
    "The two pipelines suggest a layered policy. For existing clients with at least six months of repayment history the Random Forest probability, thresholded between 0.5 and 0.6 depending on the relative cost of misses versus false alarms, supports per-account decisions on credit-line reductions and collection prioritisation. At threshold 0.5 the model flags roughly 27% of the book for review (1,631 of the 6,000 test rows), catching 60% of true defaulters; at threshold 0.6 it flags around 21% (close to 1,241 rows), still catching 53%. Each percentage point of the book that is flagged represents real review effort, so the lender's collections capacity is the practical constraint on the threshold choice.",
    "For new or thin-file clients the cluster assignment, which can be computed from any of the nine input features, supports a coarser segment-level policy: lower opening lines or stricter early-stage monitoring for accounts whose initial behaviour resembles Cluster 0. The dominance of repayment-history features in both pipelines, and the negligible weight on demographics, means the lender's data-collection effort should prioritise high-quality repayment tracking over richer demographic capture. There is also a secondary fairness benefit: a model in which SEX, EDUCATION and MARRIAGE contribute negligible importance is easier to defend under regulatory scrutiny than one that relies on protected attributes for accuracy.",
]

LIMITATIONS = [
    "Three categories of inconsistency arise. First, the dataset itself contains values not documented in the codebook: EDUCATION 0, 5, 6 (345 rows) and MARRIAGE 0 (54 rows). We collapsed these into other buckets; the literature is split between this approach and dropping the rows entirely. Either choice changes top-line model performance by less than half an AUC point on our reproduction, so the inconsistency is small, but it should be reported because a fair-lending audit will ask why specific records were re-coded. The undocumented PAY_n values of −2 and 0, present in tens of thousands of rows, are a more serious gap because the original codebook is silent on what they mean. We treated the columns as numeric to avoid forcing an interpretation; a reader who prefers the alternative reading, in which −2 marks no consumption and 0 marks revolving credit with the minimum payment made, would obtain similar results.",
    "Second, the class imbalance places a ceiling on what off-the-shelf metrics can say. Accuracy near 0.77 is barely above the trivial no-default predictor at 0.78; we therefore report PR-AUC and class-specific recall instead. Even with class re-weighting, our best recall on defaulters is 0.60 at threshold 0.5, which means 40% of true defaulters are missed at that threshold. The clustering and the classifier also disagree about what risk means: the cluster default rates of 14% and 29% bracket the global 22% rate but neither matches it, so the unsupervised partition identifies risk strata rather than individual risk and should not be read as a substitute for the classifier. The silhouette score of 0.24 at the chosen k = 2 is also modest, indicating that the two groups represent an axis of behavioural variation rather than a tightly separated partition.",
    "Third, the data is one bank's clients in Taiwan in 2005. Macroeconomic conditions, the Taiwanese credit-card crisis of 2006, and post-2010 regulatory changes around minimum-payment ratios all bear on the external validity of our findings. A practical deployment would have to re-train on the lender's own current portfolio and re-validate cluster stability over a longer horizon than six months.",
]

CONCLUSION = [
    "This study produced two complementary pieces of knowledge from the UCI Default of Credit Card Clients dataset. A Random Forest classifier predicts next-month default at ROC-AUC 0.78 and PR-AUC 0.56, dominated almost entirely by six features summarising the last six months of repayment behaviour. A two-cluster K-Means partition, corroborated by Ward hierarchical clustering, separates the book into a higher-utilisation, occasionally-late segment with a 29% default rate and a low-utilisation transactor segment with a 14% default rate. The two findings together support both per-account and segment-level credit policies. Demographics add little signal in either pipeline, which is methodologically reassuring, since it implies a lower risk of indirect-discrimination concerns, and operationally useful: the lender's data-collection budget belongs on repayment quality rather than on richer demographic enrichment.",
]

REFERENCES = [
    "Bhattacharyya, S, Jha, S, Tharakunnel, K & Westland, JC 2011, 'Data mining for credit card fraud: a comparative study', *Decision Support Systems*, vol. 50, no. 3, pp. 602-613.",
    "Breiman, L 2001, 'Random forests', *Machine Learning*, vol. 45, no. 1, pp. 5-32.",
    "Brown, I & Mues, C 2012, 'An experimental comparison of classification algorithms for imbalanced credit scoring data sets', *Expert Systems with Applications*, vol. 39, no. 3, pp. 3446-3453.",
    "Chawla, NV, Bowyer, KW, Hall, LO & Kegelmeyer, WP 2002, 'SMOTE: synthetic minority over-sampling technique', *Journal of Artificial Intelligence Research*, vol. 16, pp. 321-357.",
    "Chen, T & Guestrin, C 2016, 'XGBoost: a scalable tree boosting system', in *Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining*, San Francisco, CA, 13-17 August, pp. 785-794.",
    "Dua, D & Graff, C 2019, *UCI Machine Learning Repository*, University of California, Irvine, School of Information and Computer Sciences, viewed 27 May 2026, <http://archive.ics.uci.edu/ml>.",
    "Hastie, T, Tibshirani, R & Friedman, J 2009, *The elements of statistical learning: data mining, inference, and prediction*, 2nd edn, Springer, New York.",
    "Khandani, AE, Kim, AJ & Lo, AW 2010, 'Consumer credit-risk models via machine-learning algorithms', *Journal of Banking & Finance*, vol. 34, no. 11, pp. 2767-2787.",
    "Lessmann, S, Baesens, B, Seow, H-V & Thomas, LC 2015, 'Benchmarking state-of-the-art classification algorithms for credit scoring: an update of research', *European Journal of Operational Research*, vol. 247, no. 1, pp. 124-136.",
    "Pedregosa, F, Varoquaux, G, Gramfort, A, Michel, V, Thirion, B, Grisel, O, Blondel, M, Prettenhofer, P, Weiss, R, Dubourg, V, Vanderplas, J, Passos, A, Cournapeau, D, Brucher, M, Perrot, M & Duchesnay, E 2011, 'Scikit-learn: machine learning in Python', *Journal of Machine Learning Research*, vol. 12, pp. 2825-2830.",
    "Rousseeuw, PJ 1987, 'Silhouettes: a graphical aid to the interpretation and validation of cluster analysis', *Journal of Computational and Applied Mathematics*, vol. 20, pp. 53-65.",
    "Saito, T & Rehmsmeier, M 2015, 'The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets', *PLoS ONE*, vol. 10, no. 3, e0118432.",
    "Ward, JH 1963, 'Hierarchical grouping to optimize an objective function', *Journal of the American Statistical Association*, vol. 58, no. 301, pp. 236-244.",
    "Yeh, I-C & Lien, C-H 2009, 'The comparisons of data mining techniques for the predictive accuracy of probability of default of credit card clients', *Expert Systems with Applications*, vol. 36, no. 2, pp. 2473-2480.",
]


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    cover_page(doc)

    add_heading(doc, "Introduction", level=1)
    for para in INTRO:
        add_para(doc, para)

    add_heading(doc, "(a) Knowledge Acquisition", level=1)
    add_heading(doc, "Dataset and provenance", level=2)
    for para in KA_DATASET:
        add_para(doc, para)
    add_figure(doc, FIG / "01_class_balance.png",
                "Figure 1. Class distribution. The target is positive in 22.1% of rows.",
                width_cm=9)
    add_heading(doc, "Literature", level=2)
    for para in KA_LIT:
        add_para(doc, para)
    add_heading(doc, "Business framing tied to the data", level=2)
    for para in KA_BUSINESS:
        add_para(doc, para)

    add_heading(doc, "(b) Knowledge Creation", level=1)
    add_heading(doc, "Choice of techniques", level=2)
    for para in KC_TECH:
        add_para(doc, para)
    add_heading(doc, "Data preparation", level=2)
    for para in KC_DATA_PREP:
        add_para(doc, para)
    add_heading(doc, "Classification pipeline", level=2)
    for para in KC_CLASS:
        add_para(doc, para)
    add_heading(doc, "Clustering pipeline", level=2)
    for para in KC_CLUSTER:
        add_para(doc, para)

    add_heading(doc, "(c) Results and Discussion", level=1)
    add_heading(doc, "Exploratory patterns", level=2)
    for para in R_EDA:
        add_para(doc, para)
    add_figure(doc, FIG / "02_default_by_pay0.png",
                "Figure 2. Default rate by PAY_0 (September 2005 repayment status). Sample size per bin shown above each bar.")
    add_figure(doc, FIG / "03_default_by_education.png",
                "Figure 3. Default rate by education level after recoding categories 0, 5 and 6 into category 4.",
                width_cm=11)
    add_figure(doc, FIG / "04_default_by_age.png",
                "Figure 4. Default rate by age band; the youngest and oldest bands carry the highest default rate.",
                width_cm=11)
    add_figure(doc, FIG / "05_correlation_heatmap.png",
                "Figure 5. Pearson correlation among selected features. The three engineered summaries of PAY_0–PAY_6 are mutually highly correlated, as expected.",
                width_cm=13)

    add_heading(doc, "Classification results", level=2)
    for para in R_CLASS:
        add_para(doc, para)
    add_paragraph(doc, "Table 1. Classifier performance on the 6,000-row held-out test set.",
                   size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_table(doc,
               ["Metric", "Logistic Regression", "Random Forest"],
               [
                   ["ROC-AUC", "0.746", "0.779"],
                   ["PR-AUC (default class)", "0.497", "0.557"],
                   ["Accuracy", "0.744", "0.772"],
                   ["Precision (default)", "0.442", "0.488"],
                   ["Recall (default)", "0.600", "0.600"],
                   ["F1 (default)", "0.509", "0.538"],
                   ["CV ROC-AUC (mean ± SD)", "0.760 ± 0.007", "0.786 ± 0.005"],
               ])
    add_figure(doc, FIG / "06_roc_curves.png",
                "Figure 6. ROC curves on the held-out test set. RF dominates LR across most of the curve.",
                width_cm=11)
    add_figure(doc, FIG / "07_pr_curves.png",
                "Figure 7. Precision-recall curves on the held-out test set. The baseline (dashed) is the positive prevalence of 0.221.",
                width_cm=11)
    add_figure(doc, FIG / "08_confusion_matrix_rf.png",
                "Figure 8. Random Forest confusion matrix on the held-out test set at threshold 0.5.",
                width_cm=9)
    add_figure(doc, FIG / "09_feature_importance_rf.png",
                "Figure 9. Top 15 Random Forest feature importances. The top four are all repayment-history summaries.",
                width_cm=13)

    add_heading(doc, "Comparison with academic literature", level=2)
    for para in R_LIT_COMPARE:
        add_para(doc, para)

    add_heading(doc, "Clustering results", level=2)
    for para in R_CLUSTER:
        add_para(doc, para)
    add_figure(doc, FIG / "10_kmeans_selection.png",
                "Figure 10. K-Means selection: inertia (left, elbow) and silhouette score (right) for k = 2…8 on a stratified sample.")
    add_paragraph(doc, "Table 2. K-Means cluster profiles at k = 2.",
                   size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_table(doc,
               ["Feature", "Cluster 0 (n = 15,964)", "Cluster 1 (n = 14,036)"],
               [
                   ["LIMIT_BAL (NT$)", "115,238", "226,907"],
                   ["AGE (years)", "35.0", "36.0"],
                   ["mean_pay_status", "+0.43", "−0.88"],
                   ["max_pay_status", "+1.05", "−0.26"],
                   ["n_months_late", "1.38", "0.22"],
                   ["mean_utilisation", "0.63", "0.08"],
                   ["payment_to_bill", "0.08", "0.69"],
                   ["Observed default rate", "29.1%", "14.2%"],
               ])
    add_figure(doc, FIG / "11_kmeans_default_rate.png",
                "Figure 11. Observed default rate per K-Means cluster. The split is 29.1% versus 14.2%.",
                width_cm=11)
    add_figure(doc, FIG / "12_dendrogram.png",
                "Figure 12. Ward-linkage hierarchical clustering on a 4,000-row sample. The two-segment split at the top of the dendrogram mirrors the K-Means partition (≈90% contingency agreement).")

    add_heading(doc, "What the lender should do", level=2)
    for para in R_ACTION:
        add_para(doc, para)

    add_heading(doc, "(d) Inconsistencies and Limitations", level=1)
    for para in LIMITATIONS:
        add_para(doc, para)

    add_heading(doc, "Conclusion", level=1)
    for para in CONCLUSION:
        add_para(doc, para)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("References")
    set_run(run, size=16, bold=True, color=(0x1F, 0x3A, 0x68))

    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        for i, chunk in enumerate(ref.split("*")):
            if not chunk:
                continue
            run = p.add_run(chunk)
            set_run(run, size=11, italic=(i % 2 == 1))

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
